from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        print(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text, end=' | ')
            print()

file_path = "./file_load/fixtures/test_word.docx"
read_docx(file_path=file_path)